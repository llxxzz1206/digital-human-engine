"""P1 文档解析功能测试"""
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from app.rag.document_parser import parse_document, get_chunk_size, DocumentParseError


def test_txt_parsing():
    """测试纯文本解析"""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("这是测试文本。\n\n第二段内容，用于验证分块。")
        f.flush()
        result = parse_document(f.name)

    assert result["text"] == "这是测试文本。\n\n第二段内容，用于验证分块。"
    assert result["metadata"]["file_type"] == "txt"
    assert result["metadata"]["page_count"] is None
    print("[PASS] TXT 解析")


def test_md_parsing():
    """测试 Markdown 解析"""
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write("# 标题\n\n正文内容")
        f.flush()
        result = parse_document(f.name)

    assert "# 标题" in result["text"]
    assert result["metadata"]["file_type"] == "md"
    print("[PASS] MD 解析")


def test_pdf_parsing():
    """测试 PDF 解析（生成一个简单 PDF）"""
    import fitz

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name

    # 创建一个简单 PDF（英文，避免 CJK 字体问题）
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hospital Guide Test Document\n\nOrthopedics: Floor 3 East\nInternal Medicine: Floor 2 West")
    doc.save(pdf_path)
    doc.close()

    result = parse_document(pdf_path)
    assert "Hospital Guide" in result["text"]
    assert "Orthopedics" in result["text"]
    assert result["metadata"]["file_type"] == "pdf"
    assert result["metadata"]["page_count"] == 1
    print("[PASS] PDF 解析")

    Path(pdf_path).unlink()


def test_docx_parsing():
    """测试 Word 解析"""
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_path = f.name

    doc = Document()
    doc.add_paragraph("博物馆讲解知识库")
    doc.add_paragraph("青铜器展区位于二楼大厅")
    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "展品"
    table.cell(0, 1).text = "年代"
    table.cell(1, 0).text = "司母戊鼎"
    table.cell(1, 1).text = "商代"
    doc.save(docx_path)

    result = parse_document(docx_path)
    assert "博物馆讲解知识库" in result["text"]
    assert "青铜器展区" in result["text"]
    assert "司母戊鼎" in result["text"]  # 表格内容
    assert result["metadata"]["file_type"] == "docx"
    print("[PASS] DOCX 解析（含表格）")

    Path(docx_path).unlink()


def test_unsupported_format():
    """测试不支持的格式"""
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        f.write(b"fake")
        f.flush()
        try:
            parse_document(f.name)
            assert False, "应该抛出异常"
        except DocumentParseError as e:
            assert "不支持" in str(e)
    print("[PASS] 不支持格式拒绝")

    Path(f.name).unlink()


def test_file_not_found():
    """测试文件不存在"""
    try:
        parse_document("/nonexistent/file.pdf")
        assert False, "应该抛出异常"
    except DocumentParseError as e:
        assert "不存在" in str(e)
    print("[PASS] 文件不存在处理")


def test_chunk_size_map():
    """测试自适应分块大小"""
    assert get_chunk_size("pdf") == 800
    assert get_chunk_size("docx") == 600
    assert get_chunk_size("txt") == 500
    assert get_chunk_size("md") == 600
    assert get_chunk_size("unknown") == 500  # 默认
    print("[PASS] 自适应分块大小")


def test_split_text_paragraphs():
    """测试按段落分块"""
    from app.rag.knowledge_builder import knowledge_builder

    # 构造超过 500 字符的多段落文本
    paras = [f"第{i}段：{'内容' * 50}" for i in range(10)]
    text = "\n\n".join(paras)
    metadata = {"filename": "test.txt", "file_type": "txt"}

    chunks = knowledge_builder._split_text(text, metadata, 500)
    assert len(chunks) > 1
    # 每个 chunk 不应该超过太多（允许段落略超）
    for chunk in chunks:
        assert len(chunk["text"]) <= 600  # 段落边界可能略超
    print(f"[PASS] 段落分块: {len(text)} 字符 → {len(chunks)} 块")


if __name__ == "__main__":
    print("=" * 50)
    print("P1 文档解析功能测试")
    print("=" * 50)

    test_txt_parsing()
    test_md_parsing()
    test_pdf_parsing()
    test_docx_parsing()
    test_unsupported_format()
    test_file_not_found()
    test_chunk_size_map()
    test_split_text_paragraphs()

    print("=" * 50)
    print("全部通过!")

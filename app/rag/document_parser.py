"""文档解析器：支持 PDF / Word / 纯文本提取"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

# 各类型文档的推荐分块大小（字符数）
CHUNK_SIZE_MAP = {
    ".pdf": 800,
    ".docx": 600,
    ".doc": 600,
    ".txt": 500,
    ".md": 600,
}


class DocumentParseError(Exception):
    """文档解析失败"""


def parse_document(file_path: str | Path) -> dict[str, Any]:
    """解析文档，提取纯文本内容

    Args:
        file_path: 文件路径

    Returns:
        {"text": str, "metadata": {"filename": str, "file_type": str, "page_count": int|None}}

    Raises:
        DocumentParseError: 文件不存在、格式不支持或解析失败
    """
    path = Path(file_path)

    if not path.exists():
        raise DocumentParseError(f"文件不存在: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError(
            f"不支持的文件格式: {suffix}，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        if suffix == ".pdf":
            return _parse_pdf(path)
        elif suffix in (".docx", ".doc"):
            return _parse_word(path)
        else:
            return _parse_text(path)
    except DocumentParseError:
        raise
    except Exception as e:
        raise DocumentParseError(f"解析失败 [{path.name}]: {e}") from e


def _parse_pdf(path: Path) -> dict[str, Any]:
    """使用 PyMuPDF 解析 PDF"""
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    pages_text: list[str] = []

    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text.strip())

    doc.close()

    if not pages_text:
        raise DocumentParseError(f"PDF 无可提取文本（可能是扫描件）: {path.name}")

    full_text = "\n\n".join(pages_text)
    logger.info("PDF 解析完成: %s, %d 页, %d 字符", path.name, len(pages_text), len(full_text))

    return {
        "text": full_text,
        "metadata": {
            "filename": path.name,
            "file_type": "pdf",
            "page_count": len(pages_text),
        },
    }


def _parse_word(path: Path) -> dict[str, Any]:
    """使用 python-docx 解析 Word 文档"""
    if path.suffix.lower() == ".doc":
        raise DocumentParseError(
            f"不支持旧版 .doc 格式，请转换为 .docx 后上传: {path.name}"
        )

    from docx import Document

    doc = Document(str(path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise DocumentParseError(f"Word 文档无可提取文本: {path.name}")

    full_text = "\n\n".join(paragraphs)
    logger.info("Word 解析完成: %s, %d 段落, %d 字符", path.name, len(paragraphs), len(full_text))

    return {
        "text": full_text,
        "metadata": {
            "filename": path.name,
            "file_type": "docx",
            "page_count": None,
        },
    }


def _parse_text(path: Path) -> dict[str, Any]:
    """解析纯文本 / Markdown 文件"""
    encoding_candidates = ["utf-8", "gbk", "gb2312", "latin-1"]
    text = None

    for enc in encoding_candidates:
        try:
            text = path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue

    if text is None:
        raise DocumentParseError(f"无法识别文件编码: {path.name}")

    if not text.strip():
        raise DocumentParseError(f"文件内容为空: {path.name}")

    logger.info("文本解析完成: %s, %d 字符", path.name, len(text))

    return {
        "text": text.strip(),
        "metadata": {
            "filename": path.name,
            "file_type": path.suffix.lstrip("."),
            "page_count": None,
        },
    }


def get_chunk_size(file_type: str) -> int:
    """根据文件类型返回推荐分块大小"""
    return CHUNK_SIZE_MAP.get(f".{file_type}", 500)
